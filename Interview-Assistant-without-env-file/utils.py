import PyPDF2
import io
import docx
from langchain.docstore.document import Document
import streamlit as st

@st.cache_data
def extract_text(file):
    print(file)
    fileType = file.type
    content = ""
    if fileType == "application/pdf":
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file.read()))
        num_pages = len(pdf_reader.pages)
        for i in range(num_pages):
            page = pdf_reader.pages[i]
            content += page.extract_text()
    elif fileType == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = docx.Document(io.BytesIO(file.read()))
        for para in doc.paragraphs:
            content += para.text
            content += "\n"
    return content

@st.cache_data
def get_langchain_doc(file):
    metadata = {"file": file}
    doc = Document(page_content=file, metadata=metadata)
    return doc

def get_resume_initial_message(text):
    messages = [{"role": "system",
                 "content": """As an AI Hiring specialist, analyze the parsed text from a resume and generate a
                  potential interview question based on the candidate's skills, experience, and qualifications? 
                 Make sure you are not asking inappropriate question.Also make sure that when you are asking questions about candidate
                 work experience, explicitly mention the company name along with your question.
                 
                 Resume:{doc}"""},

                {"role": "assistant",
                 "content": "How did you acquire new sales revenue of 36% at United First Financial?"},
                {"role": "assistant",
                 "content": "How did you acquire customer retention rate by 36% at PWC?"},


                {"role": "user", "content": text}]
    return messages

def get_resume_actionable_insights_initial_message(text):
    messages = [{"role": "system",
                 "content": """As an AI Hiring specialist, analyze the parsed text from a resume and generate a list 
                     of potential interview questions which have actionable insights like increased revenue of client by 50% yoy based on the candidate's work experience ? 
                     Also make sure that when you are asking questions about candidate
                     work experience, always include the candidate's company name along with your question.
                    Only generate questions that includes some numbers or quantitative metrics to be measured. To be specific
                    only ask questions involving number, percentages or metrics"""},

                {"role": "assistant",
                 "content": "How did you acquire new sales revenue of 36% at United First Financial?"},
                {"role": "assistant",
                 "content": '''How did you Increased satisfaction levels by reducing the time to deployment by 70% using
                    automation as the Head of Global Support at Reynen Corporation'''},

                {"role":"assistant",
                 "content":"How did you improved customer satisfaction leel by 5% at Zycus Infotech"},

                {"role": "user", "content": text}]
    return messages

def get_user_work_exp(text):
    messages = [{"role": "system",
                 "content": """As an AI agent, your task is to analyze the parsed text and extract the work experiences of the 
                 candidate as seperate element of a list. the output should be in the form of a python list. Suppose a candida has worked 
                 in three different organizations then there will be three elemnets in the list and each ekement should contain details about the work done 
                 by the candidate in that organization"""},

                {"role": "user", "content": text}]
    return messages




def get_jd_initial_message(text):
    messages = [{"role": "system",
                 "content": """As an AI Hiring specialist, analyze the job description and generate a list of 
                 potential interview questions based on the relevant skills and requirements. Make sure you do not 
                 ask inappropriate questions."""},
                {"role": "user", "content": text}]
    return messages


def get_resume_and_JD_initial_message(text):
    messages = [{"role": "system",
                 "content": """As an AI Hiring specialist, utilize the provided resume and job description as inputs 
                 to generate a set of interview questions that align with the candidate's skills and qualifications. 
                 Ensure that the questions are relevant and appropriate, avoiding any inappropriate inquiries during 
                 the interview process."""},
                {"role": "user", "content": text}]
    return messages


def get_audio_initial_message(text):
    messages = [{"role": "system",
                 "content": """As an AI Hiring specialist, analyze the candidate response and generate a list of 
                 potential follow-up interview questions based on transcription of cadidates audio response. Make 
                 sure you do not ask inappropriate questions."""},
                {"role": "user", "content": text}]
    return messages


def get_resume_review_initial_message(text):
    messages = [{"role": "system",
                 "content": """"I want you to act as a skilled resume reviewer and evaluate my resume to help me 
                 achieve my career goals. Provide comprehensive feedback on the overall format, layout, and design of 
                 my resume to ensure it is visually appealing and professional. Analyze the content of each section, 
                 including the summary, work experience, education, and skills, and offer valuable suggestions to 
                 highlight my key achievements and contributions effectively. Additionally, assess the use of 
                 keywords and industry-specific terminology to tailor my resume for specific job opportunities."""},
                {"role": "user", "content": text}]
    return messages

def get_conversation_string():
    conversation_string = ""
    for i in range(len(st.session_state['responses'])-1):
        
        conversation_string += "Human: "+st.session_state['requests'][i] + "\n"
        conversation_string += "Bot: "+ st.session_state['responses'][i+1] + "\n"
    return conversation_string
